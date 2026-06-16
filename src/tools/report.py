"""报告生成工具 - 基于 FPDF2 (安全版: 不暴露服务器路径)"""

import json
import logging
import os
import uuid
import re
from datetime import datetime

logger = logging.getLogger(__name__)

REPORTS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "data", "reports")


def _clean_markdown(content: str) -> str:
    """清理 Markdown 内容"""
    # 清理 HTML 标签
    content = re.sub(r'<[^>]+>', '', content)
    # 清理多余空行
    content = re.sub(r'\n{3,}', '\n\n', content)
    return content.strip()


def generate_report_impl(content: str, title: str, format: str = "pdf") -> str:
    """生成 PDF 或 DOCX 报告，返回 download_url（不暴露服务器路径）"""
    try:
        os.makedirs(REPORTS_DIR, exist_ok=True)
        content = _clean_markdown(content)

        # 修复日期 — 只替换"报告生成时间"标注行，不修改正文历史数据
        current_date = datetime.now().strftime("%Y年%m月%d日")
        content = re.sub(r'报告生成时间[:：]\s*\d{4}年\d{1,2}月\d{1,2}日', f'报告生成时间：{current_date}', content)
        content = re.sub(r'报告生成时间[:：]\s*\d{4}年', f'报告生成时间：{datetime.now().year}年', content)

        safe_title = re.sub(r'[^\w\-]', '_', title)
        report_id = f"{safe_title}_{uuid.uuid4().hex[:8]}"

        if format.lower() == "docx":
            filepath = _generate_docx(content, report_id)
        else:
            filepath = _generate_pdf(content, report_id)

        # Return download_url instead of raw file_path (security fix)
        ext = ".docx" if format.lower() == "docx" else ".pdf"
        download_url = f"/download/{report_id}{ext}"

        return json.dumps({
            "status": "success",
            "message": f"报告已生成",
            "download_url": download_url,
            "format": format,
            "title": title,
            "report_id": report_id,
        }, ensure_ascii=False)

    except Exception as e:
        logger.error(f"报告生成失败: {e}")
        return json.dumps({
            "status": "error",
            "message": f"报告生成失败: {str(e)[:200]}"
        }, ensure_ascii=False)


def _get_report_filepath(report_id: str) -> str | None:
    """根据 report_id 查找实际文件路径（供 download 端点使用）"""
    for ext in [".pdf", ".docx"]:
        path = os.path.join(REPORTS_DIR, f"{report_id}{ext}")
        if os.path.exists(path):
            return path
    return None


def _generate_pdf(content: str, report_id: str) -> str:
    """使用 FPDF2 生成 PDF"""
    from fpdf import FPDF

    filepath = os.path.join(REPORTS_DIR, f"{report_id}.pdf")

    pdf = FPDF()
    pdf.add_page()

    # 注册中文字体
    font_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "assets", "NotoSansSC-Regular.ttf")
    if os.path.exists(font_path):
        pdf.add_font("CJK", "", font_path, uni=True)
        pdf.add_font("CJK", "B", font_path, uni=True)
        font_name = "CJK"
    else:
        # Fallback: 使用内置字体（不支持中文）
        font_name = "Helvetica"
        logger.warning("中文字体未找到，PDF 将无法正确渲染中文")

    pdf.set_font(font_name, "", 10)
    pdf.set_auto_page_break(auto=True, margin=15)

    # 写入内容
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            pdf.ln(5)
            continue

        # 标题处理
        if line.startswith('# ') and font_name == "CJK":
            pdf.set_font(font_name, "B", 16)
            pdf.ln(4)
            pdf.cell(0, 10, line[2:], ln=True)
            pdf.ln(2)
            pdf.set_font(font_name, "", 10)
        elif line.startswith('## ') and font_name == "CJK":
            pdf.set_font(font_name, "B", 13)
            pdf.ln(3)
            pdf.cell(0, 8, line[3:], ln=True)
            pdf.ln(1)
            pdf.set_font(font_name, "", 10)
        elif line.startswith('### ') and font_name == "CJK":
            pdf.set_font(font_name, "B", 11)
            pdf.cell(0, 7, line[4:], ln=True)
            pdf.set_font(font_name, "", 10)
        elif line.startswith('- ') or line.startswith('* '):
            pdf.cell(5)
            pdf.cell(0, 5, "  " + line[2:], ln=True)
        elif line.startswith('|') and '|' in line[1:]:
            # 简单表格处理
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if all(c.startswith('---') for c in cells):
                continue
            for i, cell in enumerate(cells):
                pdf.cell(35, 6, cell[:20], border=1)
            pdf.ln()
        else:
            pdf.multi_cell(0, 5, line)

    pdf.output(filepath)
    return filepath


def _generate_docx(content: str, report_id: str) -> str:
    """使用 python-docx 生成 DOCX"""
    from docx import Document

    filepath = os.path.join(REPORTS_DIR, f"{report_id}.docx")

    doc = Document()

    for line in content.split('\n'):
        line = line.strip()
        if not line:
            continue

        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)

    doc.save(filepath)
    return filepath
